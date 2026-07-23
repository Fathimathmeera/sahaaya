from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

ROOT = Path(r"C:\Users\QS\Documents\sahaya")
TEMPLATE = Path(r"C:\Users\QS\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.0\skills\artifact-template-design-report\assets\reference.docx")
OUT_DIR = ROOT / "output"
OUT_DIR.mkdir(exist_ok=True)
DOCX_OUT = OUT_DIR / "sahaya-design-report.docx"
PDF_OUT = OUT_DIR / "sahaya-design-report.pdf"

copy2(TEMPLATE, DOCX_OUT)
doc = Document(DOCX_OUT)

content = {
    2: "Sahaya Accessibility Platform",
    4: "Design Report",
    6: "Executive summary",
    7: "Sahaya is a Next.js accessibility companion designed around wheelchair users' everyday planning, navigation, reporting, and emergency-support needs. The current prototype combines accessible route planning, indoor navigation, building-image analysis, community reports, emergency SOS, and an accessibility-focused assistant.",
    8: "The design prioritizes clear actions, mobile-friendly controls, server-side integrations, and cautious communication. It is a college-project prototype: users should verify physical accessibility conditions and use local emergency services for immediate danger.",
    9: "At a glance",
    10: "Six feature areas are available from the authenticated Sahaya dashboard: accessible routes, indoor navigation, scanner, community reports, emergency SOS, and health/assistant support.",
    11: "Firebase supports authentication and data storage; OpenStreetMap/Leaflet, GraphHopper, Twilio, Google Places, and OpenAI are kept behind server-side routes where applicable.",
    12: "The production build completed successfully after source files were normalized to UTF-8 and the generated build cache was recreated.",
    13: "Introduction",
    14: "Wheelchair users often need more than a shortest-path map: they need context about ramps, lifts, entrances, washrooms, obstacles, and local reports. Sahaya groups these needs into one accessible web application with a concise, feature-led dashboard.",
    15: "The project deliberately separates deterministic systems from AI. Maps, routing, GPS, authentication, Firebase storage, and SMS delivery remain conventional service integrations. AI is reserved for language and image interpretation that benefits from reasoning.",
    16: "Key findings",
    17: "The current implementation has a clear client/server boundary. Browser pages call local API routes; server routes hold service credentials and return focused response data. This reduces accidental key exposure and makes each feature easier to test independently.",
    18: "Context and conditions",
    19: "Authentication is the entry point to application features. The login screen is limited to credentials and recovery links, while protected feature pages redirect unauthenticated visitors to login. The root page becomes the authenticated feature dashboard after sign-in.",
    20: "Patterns in the evidence",
    21: "Accessibility work is strongest when information is concise and qualified. The scanner reports only visible evidence, community reports retain the user's meaning, and emergency guidance avoids claiming that an exit is safe or available without verified data.",
    23: "Key takeaway. Sahaya is most credible when it helps users interpret evidence and plan next actions without presenting uncertain accessibility information as a guarantee.",
    25: "Implications",
    26: "A production deployment should treat data quality as a shared responsibility. Community reports need moderation, building scans need visible-evidence limitations, and route recommendations need local verification. The interface should make those limits easy to understand without making the product feel alarmist.",
    27: "Recommendations",
    28: "Prioritize a small, reliable demonstration path: authenticate, plan a route, upload a clear entrance photo, submit a structured community report, and request assistant guidance. This gives reviewers a complete story while keeping external configuration manageable.",
    29: "Configure production services. Add Firebase, GraphHopper, Twilio, Google Places, and OpenAI values as server environment variables. Keep OPENAI_API_KEY in .env.local locally and Vercel server settings in deployment.",
    30: "Validate with realistic scenarios. Test keyboard navigation, mobile layout, Malayalam text, denied location permissions, missing integrations, invalid image uploads, and emergency fallback messages before a demo.",
    31: "Manage AI usage deliberately. Use the Responses API only after a user requests assistant help, scans an image, prepares a report, starts indoor guidance, or asks for emergency wording. Keep short inputs, structured responses, and local fallbacks.",
    32: "Conclusion",
    33: "The Sahaya prototype has a coherent service boundary and a feature set that maps to high-impact wheelchair-accessibility tasks. Its strongest next step is field validation with representative users and local accessibility stakeholders.",
    34: "The platform should continue to favor practical, plain-language support over overconfident automation. This approach supports both accessibility and trust.",
    35: "Appendix",
    36: "Implementation notes",
    37: "Client routes include login, registration, navigation, indoor navigation, scanner, community reports, emergency SOS, and the assistant. Feature routes are guarded through the shared authentication shell.",
    38: "OpenAI is server-only. Assistant and scanner routes use the Responses API; community, indoor, and emergency routes request compact structured text only when AI interpretation adds value.",
    39: "Source notes",
    40: "Sahaya project source code and configuration, reviewed July 2026.",
    41: "Sahaya README: environment variables, deployment notes, and OpenAI cost controls.",
    42: "Next.js production build output: compiled successfully, type validation passed, and 22 application routes generated.",
    43: "OpenAI developer documentation: Responses API text and vision guidance.",
    45: "Report scope. This design report is based on the current project implementation and build verification. It does not certify physical accessibility, regulatory compliance, or live third-party service availability.",
}

for index, text in content.items():
    paragraph = doc.paragraphs[index]
    paragraph.text = text
    if index == 2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Helvetica Neue"
            run.font.size = Pt(28)
            run.font.color.rgb = RGBColor(32, 70, 92)

subtitle = doc.tables[0]
subtitle.cell(0, 0).text = "A production-readiness design review for an AI-powered wheelchair accessibility companion"
subtitle.cell(0, 2).text = "Prepared for Sahaya\nJuly 2026"

findings = doc.tables[1]
data = [
    ("Access", "Protected routes and focused dashboard", "Keep authentication simple and visible"),
    ("Evidence", "Scanner and reports qualify uncertain conditions", "Avoid unsupported accessibility claims"),
    ("AI", "Server-only Responses API usage", "Reserve calls for interpretation and language"),
]
for row, values in zip(findings.rows[1:], data):
    for cell, value in zip(row.cells, values):
        cell.text = value

doc.core_properties.title = "Sahaya Accessibility Platform - Design Report"
doc.core_properties.author = "Sahaya Project"
doc.save(DOCX_OUT)

styles = getSampleStyleSheet()
title = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=30, textColor=colors.HexColor("#20465C"), alignment=TA_CENTER, spaceAfter=14)
subtitle_style = ParagraphStyle("ReportSubtitle", parent=styles["BodyText"], fontName="Helvetica", fontSize=11, leading=15, textColor=colors.HexColor("#53616D"), alignment=TA_CENTER, spaceAfter=32)
h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#20465C"), spaceBefore=14, spaceAfter=8)
h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#39726A"), spaceBefore=10, spaceAfter=5)
body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, spaceAfter=8, textColor=colors.HexColor("#27323A"))
story = [Paragraph("Sahaya Accessibility Platform", title), Paragraph("Design Report | July 2026", subtitle_style), Paragraph("Executive summary", h1), Paragraph(content[7], body), Paragraph(content[8], body), Paragraph("At a glance", h2), Paragraph(content[10], body), Paragraph(content[11], body), Paragraph(content[12], body), Paragraph("Introduction", h2), Paragraph(content[14], body), Paragraph(content[15], body), PageBreak(), Paragraph("Key findings", h1), Paragraph(content[17], body), Paragraph("Context and conditions", h2), Paragraph(content[19], body), Paragraph("Patterns in the evidence", h2), Paragraph(content[21], body), Paragraph(content[23], body)]
table_data = [[Paragraph("Theme", body), Paragraph("Observation", body), Paragraph("Implication", body)]] + [[Paragraph(a, body), Paragraph(b, body), Paragraph(c, body)] for a,b,c in data]
table = Table(table_data, colWidths=[1.2*inch, 2.55*inch, 2.25*inch])
table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#20465C")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#C7D4D9")), ("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 8), ("RIGHTPADDING", (0,0), (-1,-1), 8), ("TOPPADDING", (0,0), (-1,-1), 7), ("BOTTOMPADDING", (0,0), (-1,-1), 7)]))
story += [Spacer(1, 8), table, Paragraph("Implications", h2), Paragraph(content[26], body), Paragraph("Recommendations", h1), Paragraph(content[28], body), Paragraph("1. Configure production services.", h2), Paragraph(content[29], body), Paragraph("2. Validate with realistic scenarios.", h2), Paragraph(content[30], body), Paragraph("3. Manage AI usage deliberately.", h2), Paragraph(content[31], body), Paragraph("Conclusion", h2), Paragraph(content[33], body), Paragraph(content[34], body), PageBreak(), Paragraph("Appendix", h1), Paragraph("Implementation notes", h2), Paragraph(content[37], body), Paragraph(content[38], body), Paragraph("Source notes", h2), Paragraph(content[40], body), Paragraph(content[41], body), Paragraph(content[42], body), Paragraph(content[43], body), Paragraph(content[45], body)]

def footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#C7D4D9"))
    canvas.line(0.75*inch, 0.55*inch, 7.75*inch, 0.55*inch)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#53616D"))
    canvas.drawString(0.75*inch, 0.38*inch, "Sahaya Design Report")
    canvas.drawRightString(7.75*inch, 0.38*inch, f"Page {doc.page}")
    canvas.restoreState()

pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch, topMargin=0.8*inch, bottomMargin=0.8*inch)
pdf.build(story, onFirstPage=footer, onLaterPages=footer)
print(DOCX_OUT)
print(PDF_OUT)
